from __future__ import annotations

import asyncio
import os
import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency fallback
    Document = None


@dataclass(frozen=True)
class SarvamTransformResult:
    text: str
    language_code: str
    instruction_applied: bool
    provider: str
    fallback: bool
    notes: tuple[str, ...]


class SarvamService:
    def __init__(self) -> None:
        self._api_key = os.getenv("SARVAM_API_KEY", "").strip()
        self._default_target = os.getenv("SARVAM_DEFAULT_TARGET_LANGUAGE", "en-IN").strip() or "en-IN"
        self._default_gender = os.getenv("SARVAM_SPEAKER_GENDER", "Male").strip() or "Male"

    def enabled(self) -> bool:
        return bool(self._api_key)

    async def transform_report(
        self,
        text: str,
        target_language_code: str,
        instruction: str | None = None,
    ) -> SarvamTransformResult:
        cleaned_text = text.strip()
        if not cleaned_text:
            return SarvamTransformResult(
                text="No report content available.",
                language_code=target_language_code or self._default_target,
                instruction_applied=bool(instruction and instruction.strip()),
                provider="none",
                fallback=True,
                notes=("Report content was empty.",),
            )

        target = (target_language_code or self._default_target).strip()
        normalized_instruction = (instruction or "").strip()

        if not self.enabled():
            transformed = _apply_instruction_fallback(cleaned_text, normalized_instruction)
            return SarvamTransformResult(
                text=transformed,
                language_code=target,
                instruction_applied=bool(normalized_instruction),
                provider="fallback",
                fallback=True,
                notes=("SARVAM_API_KEY not configured; returned fallback transformation.",),
            )

        try:
            translated = await asyncio.to_thread(
                self._translate_with_sdk,
                cleaned_text,
                target,
            )
            transformed = _apply_instruction_fallback(translated, normalized_instruction)
            return SarvamTransformResult(
                text=transformed,
                language_code=target,
                instruction_applied=bool(normalized_instruction),
                provider="sarvam",
                fallback=False,
                notes=(),
            )
        except Exception as exc:
            transformed = _apply_instruction_fallback(cleaned_text, normalized_instruction)
            error_note = _trim_error_message(exc)
            return SarvamTransformResult(
                text=transformed,
                language_code=target,
                instruction_applied=bool(normalized_instruction),
                provider="sarvam-fallback",
                fallback=True,
                notes=(f"Sarvam translation failed: {type(exc).__name__}: {error_note}",),
            )

    async def build_docx(
        self,
        title: str,
        text: str,
        target_language_code: str,
        instruction: str | None = None,
    ) -> tuple[bytes, SarvamTransformResult]:
        transformed = await self.transform_report(text, target_language_code, instruction)
        docx_bytes = _to_docx_bytes(title=title, body=transformed.text)
        return docx_bytes, transformed

    def _translate_with_sdk(self, text: str, target_language_code: str) -> str:
        from sarvamai import SarvamAI

        client = SarvamAI(api_subscription_key=self._api_key)
        # Translate in chunks to avoid provider-side payload limits on long reports.
        chunks = _chunk_text_for_translation(text)
        translated_chunks: list[str] = []
        for chunk in chunks:
            response = client.text.translate(
                input=chunk,
                source_language_code="auto",
                target_language_code=target_language_code,
            )
            resolved = _extract_translation_text(response)
            if not resolved:
                raise RuntimeError("Sarvam translation returned empty text")
            translated_chunks.append(resolved)
        return "\n\n".join(translated_chunks).strip()


def _extract_translation_text(response: Any) -> str:
    candidates: list[str] = []

    if isinstance(response, str):
        candidates.append(response)
    if isinstance(response, dict):
        for key in (
            "translated_text",
            "translation",
            "text",
            "output",
            "result",
        ):
            value = response.get(key)
            if isinstance(value, str):
                candidates.append(value)

    for attr in (
        "translated_text",
        "translation",
        "text",
        "output",
        "result",
    ):
        value = getattr(response, attr, None)
        if isinstance(value, str):
            candidates.append(value)

    for candidate in candidates:
        if candidate and candidate.strip():
            return candidate.strip()
    return ""


def _apply_instruction_fallback(text: str, instruction: str) -> str:
    if not instruction:
        return text

    lower = instruction.lower()
    compact = text

    if "short" in lower or "concise" in lower or "summar" in lower:
        compact = _summarize_light(compact)

    if "bullet" in lower:
        compact = _bullets_from_paragraphs(compact)

    if "formal" in lower:
        compact = compact.replace("can't", "cannot").replace("don't", "do not")

    header = f"Applied user instruction: {instruction}\n\n"
    return header + compact


def _summarize_light(text: str, max_sentences: int = 8) -> str:
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    selected = [s.strip() for s in sentences if s.strip()][:max_sentences]
    return " ".join(selected) if selected else text


def _bullets_from_paragraphs(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return text
    if all(line.startswith("-") for line in lines[: min(4, len(lines))]):
        return text
    return "\n".join(f"- {line}" for line in lines)


def _to_docx_bytes(title: str, body: str) -> bytes:
    if Document is None:
        fallback = (
            f"{title.strip() or 'Hexamind Research Report'}\n\n"
            f"{body}\n\n"
            "Note: python-docx is not installed in this runtime, so this is a fallback payload."
        )
        return fallback.encode("utf-8")

    document = Document()
    document.add_heading(title.strip() or "Hexamind Research Report", level=1)

    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def docx_supported() -> bool:
    return Document is not None


def _chunk_text_for_translation(text: str, max_chars: int = 900) -> list[str]:
    cleaned = text.strip()
    if len(cleaned) <= max_chars:
        return [cleaned]

    paragraphs = [segment.strip() for segment in cleaned.split("\n\n") if segment.strip()]
    if not paragraphs:
        return [cleaned]

    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        candidate = para if not current else f"{current}\n\n{para}"
        if len(candidate) <= max_chars:
            current = candidate
            continue

        if current:
            chunks.append(current)
            current = ""

        if len(para) <= max_chars:
            current = para
            continue

        # Split oversized paragraph into sentence-like fragments.
        parts = re.split(r"(?<=[.!?])\s+", para)
        buffer = ""
        for part in parts:
            part = part.strip()
            if not part:
                continue
            candidate_part = part if not buffer else f"{buffer} {part}"
            if len(candidate_part) <= max_chars:
                buffer = candidate_part
                continue
            if buffer:
                chunks.append(buffer)
            buffer = part
        if buffer:
            current = buffer

    if current:
        chunks.append(current)

    return chunks or [cleaned]


def _trim_error_message(exc: Exception, limit: int = 180) -> str:
    message = str(exc).strip() or "no error details"
    if len(message) <= limit:
        return message
    return message[: limit - 3].rstrip() + "..."
